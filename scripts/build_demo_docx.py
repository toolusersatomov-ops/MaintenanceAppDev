"""Builds /app/tare_Demo_Test_Scripts.docx - step-by-step demo scripts for every scenario."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BEET = RGBColor(0x8E, 0x24, 0x44)
INK = RGBColor(0x2A, 0x26, 0x22)
GREY = RGBColor(0x6B, 0x66, 0x60)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.color.rgb = BEET
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.color.rgb = INK
    return p


def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def step_table(rows):
    """rows: list of (step_no, action, expected)"""
    t = doc.add_table(rows=1, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text = "Step"
    hdr[1].text = "Action"
    hdr[2].text = "Expected Result"
    for c, w in zip(hdr, (Inches(0.5), Inches(3.4), Inches(2.9))):
        c.width = w
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for no, action, expected in rows:
        cells = t.add_row().cells
        cells[0].text = str(no)
        cells[1].text = action
        cells[2].text = expected
        cells[0].width = Inches(0.5)
        cells[1].width = Inches(3.4)
        cells[2].width = Inches(2.9)
    doc.add_paragraph()


# ---------------------------------------------------------------- Title page
title = doc.add_heading("tare — Demo & Test Scripts", level=0)
for r in title.runs:
    r.font.color.rgb = BEET
para("Automated Drink Vending Machine — Maintenance & Operations App", italic=True, color=GREY)
para("App URL: https://hulk-maintenance-app.preview.emergentagent.com", bold=True)
para("All passwords are 1234 — or simply click a quick-login chip on the Login page to auto-fill credentials.")
para("Before every live demo: log in as admin01 and click \u201cReset Demo Data\u201d on the Admin Dashboard. "
     "This restores every scenario in this document to its exact starting state.", bold=True)

h2("Test Accounts")
t = doc.add_table(rows=1, cols=4)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
for i, v in enumerate(["User ID", "Role", "Name", "Scope"]):
    hdr[i].text = v
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
for row in [
    ("operations_sup01", "Operations Supervisor", "Priya Sharma", "All machines"),
    ("kitchen01", "Kitchen Staff", "Rakesh Kumar", "Kitchen"),
    ("operations01", "Operations Staff", "Anil Verma", "M001, M002, M003"),
    ("operations02", "Operations Staff", "Suresh Reddy", "M004, M005"),
    ("tech01", "Maintenance Technician", "Vikram Singh", "Work orders"),
    ("maintenance_sup01", "Maintenance Supervisor", "Ravi Patel", "Maintenance"),
    ("admin01", "Admin", "System Admin", "Demo Reset / data status"),
]:
    cells = t.add_row().cells
    for i, v in enumerate(row):
        cells[i].text = v
doc.add_paragraph()
doc.add_page_break()

# ---------------------------------------------------------------- Scenario 0
h1("Scenario 0 — Login Security & Quick Chips (2 min)")
para("Shows the role-based login, quick demo chips, 5-attempt lockout and supervisor unlock.")
step_table([
    (1, "Open the app. On the Login page, click the chip \u201ckitchen01 / Kitchen Staff\u201d.",
        "User ID and Password fields auto-fill."),
    (2, "Clear the password, type a WRONG password (e.g. 9999) and click Sign In. Repeat 5 times.",
        "Each attempt shows an invalid-credentials error. After the 5th failure the account is locked and the error says the account is locked."),
    (3, "Click the chip \u201coperations_sup01\u201d and Sign In.",
        "Supervisor dashboard opens (sidebar shows \u201cPriya Sharma \u2013 Operations Supervisor\u201d)."),
    (4, "Go to User & Access Management. Find kitchen01 and click Unlock.",
        "kitchen01 status returns to Active, failed attempts reset to 0."),
    (5, "Log out. Click the \u201ckitchen01\u201d chip and Sign In with the correct password.",
        "Kitchen dashboard opens normally."),
])

# ---------------------------------------------------------------- Scenario 1
h1("Scenario 1 — Alert \u2192 Replacement, Full End-to-End (8\u201310 min)")
para("The core story: a Low Stock alert travels Supervisor \u2192 Kitchen \u2192 Operations Staff \u2192 back to Kitchen. "
     "Four logins are used; keep two browser tabs open to switch faster.")

h2("Part A — Operations Supervisor (operations_sup01)")
step_table([
    (1, "Open Dashboard.", "KPI cards show 5 machines, today\u2019s sales, cups sold, pending kitchen/ops counts."),
    (2, "Open Alerts. Pick a Low Stock alert (e.g. M002 \u2013 Almond Milk) and click it.",
        "Alert Detail page shows current quantity, level %, capacity, affected recipes and suggested action."),
    (3, "In the alert detail, choose Operations Staff = operations01 and click Assign / Create Tasks.",
        "Success message. Alert status becomes Assigned. A Kitchen Fill Ticket + Pickup Task + Bin Replacement Task are auto-created and linked."),
    (4, "Open Live Task Progress.",
        "A new timeline card for the machine shows: Alert Reviewed by Supervisor \u2192 Operations Staff Assigned \u2192 Kitchen Fill Ticket Created."),
])

h2("Part B — Kitchen Staff (kitchen01)")
step_table([
    (5, "Open Preparation Requests \u2192 Pending tab.",
        "The new fill ticket appears with machine, ingredient and auto-calculated quantity (read-only)."),
    (6, "Click Start Preparation on the ticket.", "Ticket moves to In Progress; you are guided to Bin Filling."),
    (7, "On Bin Filling: click Scan Bin QR and enter/select the QR of a clean spare bin (e.g. QR-BIN-SPARE-LIQUID-1).",
        "Bin accepted. Quantity is pre-filled (max usage \u00d7 120) and read-only; expiry auto-set from shelf life."),
    (8, "Click Save Bin.",
        "Status becomes \u201cSaved / Ready for Pickup\u201d. The Operations pickup task unblocks. Trying to save without scanning shows the exact spec error message."),
])

h2("Part C — Operations Staff (operations01)")
step_table([
    (9, "Open Pickup List and select the machine (M002).",
        "The saved bin shows status Ready for Pickup with its QR id."),
    (10, "Click Open Camera / Scan QR and enter the bin\u2019s QR id.",
        "Item flips to Picked. Counter \u201cPicked: x / y\u201d updates."),
    (11, "Click \u201cMark All Scheduled Items Picked\u201d.",
        "Trolley status becomes Loaded; task stage advances."),
    (12, "Open Bin Replacement Tasks and open the task. Follow the guided steps in order: "
         "(a) Scan Slot QR (e.g. SLOTQR-M002-L2) \u2192 (b) Remove Old Bin \u2192 (c) Scan New Bin & Place in Machine \u2192 "
         "(d) Scan Old Bin \u2014 here you may type the literal text AUTO (lenient scan).",
        "Each step validates in order; wrong QR shows the exact spec error. After (c) the Machine Control Center slot refills to 100%. Old bin is added to Dirty Bin Return."),
    (13, "Click Complete Replacement.",
        "Task status Completed, stage Closed. Alert auto-closes. Supervisor gets a notification."),
    (14, "Open Dirty Bin Return, select the machine, and scan the old bin\u2019s QR.",
        "Status becomes \u201cReturned to Kitchen\u201d. Kitchen is notified."),
])

h2("Part D — Kitchen closes the loop (kitchen01)")
step_table([
    (15, "Open Cleaning Bins.",
        "The returned bin appears with a read-only 6-step Cleaning Guide (Returned to Kitchen \u2192 Washing \u2192 Washed \u2192 Drying \u2192 Dried \u2192 Clean)."),
    (16, "Read/follow the guide, then click \u201cMark as Cleaned \u2014 Ready for Filling\u201d (single click).",
        "All lifecycle steps complete at once. The bin returns to the spare pool (visible in Bin Storage) ready for the next fill."),
])

h2("Part E — Supervisor wrap-up (operations_sup01)")
step_table([
    (17, "Open Live Task Progress.",
        "The full timeline shows every stage with timestamp and actor (who did what, when)."),
    (18, "Open Reports \u2192 Bin Replacement.",
        "The completed replacement is listed; Activity report shows the individual actions."),
])

# ---------------------------------------------------------------- Scenario 2
h1("Scenario 2 — Bulk Pre-Schedule Replacements (5 min)")
para("Supervisor schedules many bins in one order. Items that don\u2019t need the kitchen are auto-fulfilled from spare clean bins.")
para("Pre-seeded example: an M004 bulk order already exists \u2014 Pomegranate Juice + Vanilla Protein Powder (waiting for kitchen) and Banana (auto-fulfilled, Ready for Pickup).", italic=True)
step_table([
    (1, "Login operations_sup01 \u2192 open Pre-Schedule Bulk Replacements. Select machine M005.",
        "All slots for M005 render as clickable cards grouped by Liquid / Powder / Solid / Other with level %."),
    (2, "Click 3\u20134 slot cards (or use \u201cSelect All Low Stock Items\u201d).",
        "Each card is added to the Reschedule Cart at the bottom; cart count updates."),
    (3, "In the cart, toggle \u201cKitchen Required\u201d OFF for one Solid item (e.g. Banana). Set priorities/comments as desired.",
        "Item is marked kitchen-not-required."),
    (4, "Choose Operations Staff = operations02 and click Place Bulk Order.",
        "Success toast with one bulk order id. Kitchen-required items create fill tickets; the non-kitchen item is instantly auto-fulfilled from a spare clean bin."),
    (5, "Login kitchen01 \u2192 Preparation Requests.",
        "The new tickets appear grouped under a \u201cBulk Order \u00b7 n item(s)\u201d banner."),
    (6, "Login operations02 \u2192 Pickup List \u2192 select M005.",
        "Same bulk grouping. The auto-fulfilled item is already \u201cReady for Pickup\u201d while others show \u201cPending Prep / Awaiting Kitchen\u201d."),
    (7, "(Optional) Complete the flow exactly as Scenario 1 Parts B\u2013D for one of the items.",
        "Bulk items behave identically to single tasks once created."),
])

# ---------------------------------------------------------------- Scenario 3
h1("Scenario 3 — Machine Cleaning & Sanitization (3 min)")
step_table([
    (1, "Login operations01 \u2192 Cleaning & Sanitization \u2192 select machine M001.",
        "Today\u2019s 11-step checklist opens \u2014 already 3/11 completed with photos (pre-seeded demo)."),
    (2, "On the next step, click Mark Complete WITHOUT uploading a photo.",
        "Validation error: photo is required before completing a step."),
    (3, "Upload a photo (any mock image), add a comment, Mark Complete. Repeat for remaining steps.",
        "Each step turns green with photo + comment stored."),
    (4, "Complete the final step.",
        "Task status flips to Completed; machine\u2019s last-cleaning date updates; Supervisor receives a notification."),
    (5, "Login operations_sup01 \u2192 Reports \u2192 Cleaning.",
        "Both M001 (today) and M003 (yesterday, pre-seeded) completed cleanings are listed."),
])

# ---------------------------------------------------------------- Scenario 4
h1("Scenario 4 — Kitchen Change Request (2 min)")
step_table([
    (1, "Login kitchen01 \u2192 Change Requests.",
        "Pre-seeded open request visible: quantity overfilled due to human error."),
    (2, "From Preparation Requests, open any ticket and raise a new Change Request with a reason.",
        "Request created with status Open; Supervisor is notified."),
    (3, "Login operations_sup01 \u2192 open the notification / Kitchen Preparation Status.",
        "The change request details are visible for review."),
    (4, "Back as kitchen01, click Resolve on the request.",
        "Status becomes Resolved."),
])

# ---------------------------------------------------------------- Scenario 5
h1("Scenario 5 — Door Control & Assigned Machines (2 min)")
step_table([
    (1, "Login operations01 \u2192 Assigned Machines.",
        "Cards for M001\u2013M003 show status, trolley state, last visit and cleaning status."),
    (2, "Open Door Control \u2192 select M002 \u2192 choose a door \u2192 click Open Door.",
        "Action logged with actor + timestamp; appears in the log list below."),
    (3, "Click Close Door, then Confirm Door Closed.",
        "Both actions appended to the audit log."),
])

# ---------------------------------------------------------------- Scenario 6
h1("Scenario 6 — Maintenance Workflow (4 min)")
step_table([
    (1, "Login maintenance_sup01 \u2192 Technical Alerts.",
        "3 seeded alerts: Blending Motor Overheating (High), Door Sensor Fault (Medium), Nozzle Flow Irregularity (Low)."),
    (2, "Open Work Orders. Find the unassigned \u201cDoor Sensor Fault\u201d order \u2192 Assign Technician \u2192 tech01.",
        "Order status changes; tech01 is notified."),
    (3, "Login tech01 \u2192 Assigned Work Orders \u2192 accept the new order and advance through repair stages.",
        "Each stage is stamped into the order history."),
    (4, "As tech01, raise a Spare Parts Request (e.g. Door Sensor, qty 1, with reason).",
        "Request created with status Pending Approval."),
    (5, "Login maintenance_sup01 \u2192 Spare Parts Approvals \u2192 approve the request.",
        "Status becomes Approved; inventory reflects the change."),
    (6, "Show Machine Health Center, PM Planner (2 machines overdue) and Escalations (1 open).",
        "Health scores, overdue PM rows, and the open escalation render with seeded data."),
])

# ---------------------------------------------------------------- Scenario 7
h1("Scenario 7 — Supervisor Analytics & Admin (3 min)")
step_table([
    (1, "Login operations_sup01 \u2192 Reports.",
        "Sales by machine and by drink (M002 is the top seller), replacement history, cleaning history, activity log."),
    (2, "Open Machine Control Center.",
        "Slot-level live view of all 5 machines: fill %, expiry dates, replacement due dates, statuses."),
    (3, "Login admin01 \u2192 Admin Dashboard.",
        "Data Status grid shows live collection counts (users, machines, slots, tasks, sales\u2026)."),
    (4, "Click Reset Demo Data and confirm.",
        "All workflow data wiped and reseeded, including every demo scenario above. Counts refresh. App is presentation-ready again."),
])

# ---------------------------------------------------------------- Checklist
doc.add_page_break()
h1("Quick Regression Checklist")
for item in [
    "Login works for all 7 users; lockout after 5 wrong attempts; supervisor unlock works.",
    "Alert assignment creates 3 linked records (replacement task, pickup task, kitchen ticket).",
    "Kitchen cannot save a bin without scanning the correct QR; quantity is auto-calculated and read-only.",
    "Pickup scan rejects a wrong QR with the exact spec error message.",
    "Replacement flow enforces step order; old-bin scan accepts the literal text AUTO.",
    "Completing a replacement refills the slot to 100% and closes the alert.",
    "Dirty bin flow: Returned from Machine \u2192 Returned to Kitchen \u2192 one-click \u201cMark as Cleaned\u201d \u2192 Clean / Ready for Filling (back in spare pool); guide steps shown on screen.",
    "Bulk orders group items by order id on Kitchen and Operations pages; kitchen-not-required auto-fulfils from spare bins.",
    "Cleaning checklist blocks step completion without a photo; full completion stamps the machine record.",
    "Email-Kitchen escalation on an alert is only allowed once the alert is 30+ minutes old.",
    "Admin Demo Reset restores all seed + demo scenarios and is safe to run repeatedly.",
]:
    doc.add_paragraph(item, style="List Bullet")

para("")
para("Automated backend regression: run  python3 /app/scripts/e2e_test.py  (29 steps, expected 100% pass).", bold=True)

doc.save("/app/tare_Demo_Test_Scripts.docx")
print("saved /app/tare_Demo_Test_Scripts.docx")
